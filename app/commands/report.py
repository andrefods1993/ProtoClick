from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from app.managers.database_manager import db_manager
from PIL import Image
from io import BytesIO
import base64

def save_logo_to_file(logo_data_base64):
    try:
        # Decodificar a string base64 para bytes
        logo_data_bytes = base64.b64decode(logo_data_base64)
        
        # Salvar os dados da logo em um arquivo de imagem
        image = Image.open(BytesIO(logo_data_bytes))
        
        # Salvar a imagem redimensionada no arquivo
        image.save('config/logo.png')
        
        return True
    except Exception as e:
        print(f'Erro ao salvar a logo: {e}')
        return False


def gerar_relatorio_protocolo_word(id_protocolo, file_path):
    # Conectar ao banco de dados
    db_manager.conectar()

    try:
        # Executar consulta SQL para obter os dados do protocolo
        db_manager.cursor.execute(
            """
            SELECT
                P.id,
                P.data,
                T.descricao,
                C.nome AS consultor,
                CL.razao_social AS cliente,
                S.nome AS solicitante,
                P.hora_inicio,
                P.hora_termino,
                P.relatorio AS descricao_protocolo
            FROM ProtocoloAtendimento P
            LEFT JOIN TipoAtendimento T ON P.tipo_atendimento_id = T.id
            LEFT JOIN Usuario C ON P.usuario_id = C.id
            LEFT JOIN Cliente CL ON P.cliente_id = CL.id
            LEFT JOIN SolicitanteCliente S ON P.solicitante_cliente_id = S.id
            WHERE P.id = ?
            """,
            (id_protocolo,)
        )
        protocolo = db_manager.cursor.fetchone()

        db_manager.cursor.execute(
            "SELECT mensagem FROM Aviso"
        )
        avisos = db_manager.cursor.fetchall()

        if protocolo:
            # Executar consulta SQL para obter os dados da empresa
            db_manager.cursor.execute(
                """
                SELECT
                    E.razao_social,
                    E.endereco_logradouro,
                    E.endereco_numero,
                    E.endereco_bairro,
                    CD.descricao,
                    UF.uf,
                    E.contato,
                    E.email,
                    E.logo
                FROM Empresa E
                LEFT JOIN Cidade CD ON E.cidade_id = CD.id
                LEFT JOIN UF UF ON E.uf_id = UF.id
                """
            )
            empresa = db_manager.cursor.fetchone()

            empresa_logo_data = empresa[8]  # Dados da logo da empresa

            # Salvar a logo da empresa em um arquivo de imagem
            if save_logo_to_file(empresa_logo_data):
                # Abrir o modelo1_protocolo.docx
                doc = Document('config/modelo1_protocolo.docx')


                 # Acessar a tabela no cabeçalho
                header_table = doc.sections[0].header.tables[0]
                footer_table = doc.sections[0].footer.tables[0]
                # Substituir campos do cabeçalho
                for row in header_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            
                            if '{{ID}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{ID}}', str(protocolo[0]))

                            if '{{DATA}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{DATA}}', protocolo[1])
                            
                            if '[LOGO]' in paragraph.text:
                                # Limpar a célula antes de adicionar a imagem
                                for run in paragraph.runs:
                                    run.clear()
                                # Adicionar a imagem da logo
                                run = paragraph.add_run()
                                run.add_picture('config/logo.png', width=Inches(1.0))
                
                for row in footer_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            
                            if '{{EMPRESA}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{EMPRESA}}', empresa[0])

                            if '{{LOGRADOURO}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{LOGRADOURO}}', empresa[1])

                            if '{{NUMERO}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{NUMERO}}', str(empresa[2]))

                            if '{{BAIRRO}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{BAIRRO}}', empresa[3])

                            if '{{CIDADE}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{CIDADE}}', empresa[4])

                            if '{{UF}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{UF}}', empresa[5])

                            if '{{CONTATO}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{CONTATO}}', str(empresa[6]))

                            if '{{EMAIL}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace('{{EMAIL}}', empresa[7])
                
                for paragraph in doc.paragraphs:
                    
                    if '{{TIPO_ATENDIMENTO}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{TIPO_ATENDIMENTO}}', protocolo[2])

                    if '{{CONSULTOR}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{CONSULTOR}}', protocolo[3])

                    if '{{CLIENTE}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{CLIENTE}}', protocolo[4])

                    if '{{SOLICITANTE}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{SOLICITANTE}}', protocolo[5])

                    if '{{HORA_INICIO}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{HORA_INICIO}}', protocolo[6])

                    if '{{HORA_TERMINO}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{HORA_TERMINO}}', protocolo[7])
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                
                                if '{{DESCRICAO}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{DESCRICAO}}', protocolo[8])
                                    # Dividir o texto em partes separadas pelo marcador '**'
                                    parts = paragraph.text.split('**')
                                    
                                    # Limpar o parágrafo antes de adicionar o texto formatado
                                    paragraph.clear()

                                    # Adicionar cada parte do texto ao parágrafo com formatação de negrito conforme necessário
                                    for idx, part in enumerate(parts):
                                        run = paragraph.add_run(part)
                                        if idx % 2 == 1:  # Verifica se o índice é ímpar (texto entre '**')
                                            run.bold = True
                                
                                if '{{AVISOS}}' in paragraph.text:
                                    # Concatenar todos os avisos em um único texto
                                    avisos_text = '\n'.join([aviso[0] for aviso in avisos])
                                    paragraph.text = paragraph.text.replace('{{AVISOS}}', avisos_text)


                # Salvar o documento com os dados substituídos
                doc.save(file_path)
                print('Relatório em word gerado com sucesso!')
            
            else:
                print('Erro ao salvar a logo da empresa.')

        else:
            print('Protocolo não encontrado.')

    except Exception as e:
        print(f'Erro ao gerar o relatório em word: {e}')

    finally:
        db_manager.desconectar()


def gerar_relatorio_protocolo_pdf(file_path):
    
    # Caminho do arquivo DOCX a ser convertido
    caminho_docx = 'config/protocolo.docx'

    # Caminho de saída para o arquivo PDF
    caminho_pdf = file_path

    # Converter o arquivo DOCX para PDF
    try:
        convert(caminho_docx, caminho_pdf)
        print('Arquivo PDF gerado com sucesso!')
    except Exception as e:
        print(f'Erro ao converter o arquivo em PDF: {e}')
